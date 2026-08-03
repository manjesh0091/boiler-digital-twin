"""
clusters/reports/markdown_to_docx.py — lightweight Markdown -> .docx renderer.

Converts the exact Markdown string cluster1_report.py already builds into a
native Word document via python-docx, so the .docx is generated
programmatically every run (no manual copy/paste/"Save As Word" step) while
the Markdown builder in cluster1_report.py stays the single source of
content -- this module only handles LAYOUT, not what the report says.

Supports the subset of Markdown cluster1_report.py actually uses:
  - headings: #, ##, ###
  - GFM pipe tables (| a | b |, with a |---|---| separator row)
  - bullet ("- ") and ordered ("1. ") list items, INCLUDING ones whose
    text soft-wraps across multiple source lines (no blank line between
    the wrapped lines) -- these are rejoined into one logical paragraph
    before inline formatting is applied, so a **bold** span that starts on
    one wrapped line and ends on the next still renders as one bold run
    instead of literal asterisks.
  - inline **bold**, `code`, and *italic*
  - blank-line-separated paragraphs

Anything else (nested lists, links, etc.) is not needed by this report and
is not handled.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt

INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*\n]+?\*)")
BULLET_RE = re.compile(r"^-\s+(.*)")
ORDERED_RE = re.compile(r"^\d+\.\s+.*")
FOOTNOTE_RE = re.compile(r"^\\\*\s+(.*)")


def _add_inline_runs(paragraph, text: str) -> None:
    for token in INLINE_PATTERN.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)


def _is_table_row(line: str) -> bool:
    return line.startswith("|") and line.endswith("|")


def _is_separator_row(line: str) -> bool:
    cells = [c.strip() for c in line.strip("|").split("|")]
    return all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c)


def _split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip("|").split("|")]


def _add_table(document: Document, rows: list[list[str]]) -> None:
    n_cols = len(rows[0])
    table = document.add_table(rows=0, cols=n_cols)
    table.style = "Light Grid Accent 1"
    for r_idx, row_cells in enumerate(rows):
        row = table.add_row()
        for c_idx, cell_text in enumerate(row_cells):
            if c_idx >= n_cols:
                continue
            cell = row.cells[c_idx]
            p = cell.paragraphs[0]
            _add_inline_runs(p, cell_text)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
    document.add_paragraph()


def _flush_paragraph_run(document: Document, buffer_lines: list[str]) -> None:
    """Groups a run of consecutive non-blank, non-heading, non-table lines
    into logical paragraphs/list-items, rejoining any soft-wrapped
    continuation lines, then renders each with inline formatting.
    """
    if not buffer_lines:
        return

    items: list[tuple[str | None, list[str]]] = []
    current_style: str | None = None
    current_texts: list[str] = []

    def flush_item():
        if current_texts:
            items.append((current_style, list(current_texts)))

    for line in buffer_lines:
        bm = BULLET_RE.match(line)
        fm = FOOTNOTE_RE.match(line)
        if bm:
            flush_item()
            current_style = "List Bullet"
            current_texts = [bm.group(1)]
        elif fm:
            flush_item()
            current_style = None
            current_texts = ["* " + fm.group(1)]
        elif ORDERED_RE.match(line):
            flush_item()
            current_style = None
            current_texts = [line]
        else:
            if not current_texts:
                current_style = None
            current_texts.append(line)
    flush_item()

    for style, texts in items:
        text = " ".join(t for t in texts if t)
        p = document.add_paragraph(style=style) if style else document.add_paragraph()
        _add_inline_runs(p, text)


def write_docx(markdown_text: str, out_path: Path, title: str) -> None:
    """`title` sets the document's core-properties title (Word's File >
    Info metadata) only -- the visible title comes from the markdown's own
    top-level "# " heading, so it isn't duplicated on the page.
    """
    document = Document()
    document.core_properties.title = title

    lines = [line.strip() for line in markdown_text.split("\n")]
    i = 0
    n = len(lines)
    paragraph_buffer: list[str] = []

    def flush_buffer():
        _flush_paragraph_run(document, paragraph_buffer)
        paragraph_buffer.clear()

    while i < n:
        line = lines[i]

        if not line:
            flush_buffer()
            i += 1
            continue

        if _is_table_row(line):
            flush_buffer()
            table_lines = []
            while i < n and _is_table_row(lines[i]):
                table_lines.append(lines[i])
                i += 1
            rows = [_split_row(l) for l in table_lines if not _is_separator_row(l)]
            if rows:
                _add_table(document, rows)
            continue

        if line.startswith("### "):
            flush_buffer()
            document.add_heading(line[4:], level=3)
            i += 1
            continue
        if line.startswith("## "):
            flush_buffer()
            document.add_heading(line[3:], level=2)
            i += 1
            continue
        if line.startswith("# "):
            flush_buffer()
            document.add_heading(line[2:], level=1)
            i += 1
            continue

        paragraph_buffer.append(line)
        i += 1

    flush_buffer()
    document.save(str(out_path))
